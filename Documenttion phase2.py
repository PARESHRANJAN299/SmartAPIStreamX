from docx import Document

# Load the existing Phase 2 document
doc_path = "/mnt/data/SmartAPIStreamX_Phase2.docx"
doc = Document(doc_path)

# Add heading for the new enhancements
doc.add_page_break()
doc.add_heading("Enhancements Added to Phase 2", level=1)

# Section: GPT Integration for Fix Recommendations
doc.add_heading("1. GPT Integration for Fix Recommendations", level=2)
doc.add_paragraph("""
• Integrated Azure OpenAI (GPT) API with Log Analytics error output for dynamic Root Cause Analysis (RCA) and fix suggestions.
• Logs with status == "Failed" are analyzed using GPT prompt chaining to provide contextual solutions.
• Responses are displayed in the dashboard or routed to notifications (Email, Teams).
""")
doc.add_heading("Advantages for the Organization:", level=3)
doc.add_paragraph("""
✔ Faster resolution of pipeline failures with intelligent suggestions.
✔ Reduces manual troubleshooting and accelerates recovery.
✔ Leverages organizational error history to train GPT prompts for accuracy.
""")
doc.add_heading("Key Features & How It Works:", level=3)
doc.add_paragraph("""
1. Extract failed logs using KQL in Log Analytics.
2. Send error string to Azure OpenAI using Logic App or Function App.
3. Receive RCA + Fix.
4. Use in dashboard, email, or automation logic.
""")

# Section: Alert System (Logic App + Email)
doc.add_heading("2. Alert System (Logic App + Email)", level=2)
doc.add_paragraph("""
• Azure Logic App triggers on ADF failure events (via Event Grid).
• Parses error, formats HTML alert email.
• Sends notifications via Outlook or SendGrid with actionable insights.
""")
doc.add_heading("Advantages for the Organization:", level=3)
doc.add_paragraph("""
✔ Real-time alerting with full context of failure.
✔ Improves SLA response and engineering efficiency.
✔ Customizable templates and alert thresholds.
""")
doc.add_heading("Key Features & How It Works:", level=3)
doc.add_paragraph("""
1. Logic App listens to pipeline failure event.
2. Parses ADF metadata (Pipeline, Activity, Error).
3. Formats HTML using built-in Compose and Send Email.
""")

# Section: Retry / Self-Heal Logic
doc.add_heading("3. Retry / Self-Heal Logic", level=2)
doc.add_paragraph("""
• ADF native retry policy used for transient errors with exponential backoff.
• For fatal errors, Logic App queues the re-run of failed activity via ADF REST API.
• Optional: Update watermark to resume from last known good state.
""")
doc.add_heading("Advantages for the Organization:", level=3)
doc.add_paragraph("""
✔ Increased reliability and resiliency of data pipelines.
✔ Reduces manual intervention and escalations.
✔ Enhances SLA compliance and pipeline availability.
""")
doc.add_heading("Key Features & How It Works:", level=3)
doc.add_paragraph("""
1. Configure retry policy in Copy/Lookup Activity (exponential retry).
2. Use Logic App to re-trigger ADF pipeline using stored state (watermark).
3. Integrate with monitoring dashboard for visibility.
""")

# Save updated document
updated_doc_path = "/mnt/data/SmartAPIStreamX_Phase2_Updated.docx"
doc.save(updated_doc_path)
updated_doc_path
